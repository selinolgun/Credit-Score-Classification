import streamlit as st
from docx import Document
from docx.shared import Cm
from collections.abc import Sequence  # Import Sequence from collections.abc


png_list = ["images/Missing_value.png", "images/Age_count.png", "images/Age_salary.png", "images/Occupation_salary.png",
               "images/Age_occupation_salary.png", "images/Age_occupation_salary_min.png", "images/Age_distribution.png",
               "images/Age_density.png", "images/Salary_distribution.png", "images/Salary_density.png", "images/Count_occupation.png",
               "images/Montly_salary_amount.png", "images/Heatmap.png", "images/Pairplot.png", "images/Displot_salary.png",
               "images/Displot_delay.png", "images/Displot_credit.png"]

head_list = ["Missing Values Heatmap", "Age Count", "Age Salary", "Occupation Salary",
                  "Age Occupation Salary Max", "Age Occupation Salary Min", "Age Distribution",
                  "Age Density", "Salary Distribution", "Salary Density", "Count Occupation",
                  "Montly Salary Amount", "Heatmap", "Pairplot", "Displot Salary", "Displot Delay",
                  "Displot Credit"]

st.title('Credit Score Classification')

doc = Document()

number = 0
for head, png in zip(head_list, png_list):
    st.header(head)
    st.image(png, caption=head, use_column_width=True)

    doc.add_heading(head, level=2)
    doc.add_picture(png, width=Cm(14), height=Cm(8.5))

    if number % 2 != 1:
        doc.add_paragraph().paragraph_format.space_after = 0
    number += 1

    if number % 2 == 0:
        doc.add_page_break()
